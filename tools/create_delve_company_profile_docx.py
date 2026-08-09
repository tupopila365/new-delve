from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "docx"
OUTPUT_DOCX = OUTPUT_DIR / "Delve_Company_Profile_2026.docx"
LOGO_PATH = Path(r"C:\Users\kauna\Downloads\DELVE.jpg")
ASSET_DIR = ROOT / "tmp" / "pdfs" / "delve_company_profile"

PURPLE = "8C52FF"
DEEP_PURPLE = "5F2FC9"
INK = "17141D"
MUTED = "67616F"
LILAC = "F1E9FF"
PALE = "FAF8FC"
LINE = "DED7E8"
MINT = "DDF7EF"
MINT_DARK = "19765D"
GOLD = "FFF0D5"
GOLD_DARK = "9A5B00"
WHITE = "FFFFFF"

# Base: narrative_proposal. Named overrides: A4 page, Delve purple palette,
# left-aligned body for company-profile scanning, and an editorial cover.


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_run(run, size=11, bold=False, color=INK, italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = rgb(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size=6, inside=True):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    names = ["top", "left", "bottom", "right"]
    if inside:
        names += ["insideH", "insideV"]
    for name in names:
        edge = borders.find(qn(f"w:{name}"))
        if edge is None:
            edge = OxmlElement(f"w:{name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = borders.find(qn(f"w:{name}"))
        if edge is None:
            edge = OxmlElement(f"w:{name}")
            borders.append(edge)
        edge.set(qn("w:val"), "nil")


def set_table_geometry(table, widths_inches, indent_dxa=0):
    table.autofit = False
    total_dxa = round(sum(widths_inches) * 1440)
    tbl_pr = table._tbl.tblPr

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_inches:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(round(width * 1440)))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_inches[idx]
            cell.width = Inches(width)
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(round(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def clear_cell(cell):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def add_para(
    container,
    text="",
    *,
    size=11,
    bold=False,
    color=INK,
    italic=False,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    before=0,
    after=8,
    line=1.333,
    keep_with_next=False,
):
    p = container.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep_with_next
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color, italic=italic)
    return p


def add_kicker(container, text):
    return add_para(
        container,
        text.upper(),
        size=8,
        bold=True,
        color=PURPLE,
        after=6,
        line=1.0,
        keep_with_next=True,
    )


def add_heading(container, text, level=1):
    p = container.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    sizes = {1: 23, 2: 15, 3: 11.5}
    set_run(run, size=sizes[level], bold=True, color=INK if level == 1 else PURPLE)
    return p


def add_bullet(container, text, color=INK):
    p = container.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    if p.runs:
        p.runs[0].text = text
        set_run(p.runs[0], size=9.5, color=color)
    else:
        set_run(p.add_run(text), size=9.5, color=color)
    return p


def add_numbered(container, text):
    p = container.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.208
    if p.runs:
        p.runs[0].text = text
        set_run(p.runs[0], size=9.5)
    else:
        set_run(p.add_run(text), size=9.5)
    return p


def add_callout(doc, label, text, fill=LILAC, accent=PURPLE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.95])
    remove_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = clear_cell(cell)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(label.upper())
    set_run(run, size=8, bold=True, color=accent)
    add_para(cell, text, size=10, after=2, line=1.25)
    add_para(doc, "", size=2, after=4, line=1.0)
    return table


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, value, end])
    set_run(run, size=8, color=MUTED)


def set_page_furniture(doc):
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("DELVE COMPANY PROFILE 2026"), size=7.5, bold=True, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    set_run(
        p.add_run("DELVE COMPANY PROFILE 2026  |  EXPLORE. CONNECT. BELONG."),
        size=7.5,
        bold=True,
        color=MUTED,
    )


def set_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for level, size, before, after in (
        (1, 23, 18, 10),
        (2, 15, 12, 6),
        (3, 11.5, 8, 4),
    ):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(INK if level == 1 else PURPLE)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        style = doc.styles[list_style]
        style.font.name = "Calibri"
        style.font.size = Pt(9.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208


def configure_page(doc):
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def build_cover(doc):
    add_para(
        doc,
        "COMPANY PROFILE  |  2026",
        size=9,
        bold=True,
        color=PURPLE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=16,
        line=1.0,
    )
    logo = doc.add_paragraph()
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo.paragraph_format.space_after = Pt(16)
    logo_run = logo.add_run()
    logo_run.add_picture(str(LOGO_PATH), width=Inches(5.25))
    add_para(
        doc,
        "EVERYONE IS A TRAVELLER.",
        size=20,
        bold=True,
        color=PURPLE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=6,
        line=1.0,
    )
    add_para(
        doc,
        "Explore. Connect. Belong.",
        size=12,
        bold=True,
        color=INK,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=4,
        line=1.0,
    )
    add_para(
        doc,
        "Built in Namibia for a connected world.",
        size=10,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=24,
        line=1.0,
    )
    add_para(
        doc,
        "delveworldwide@gmail.com  |  +264 81 764 9719",
        size=9,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=0,
        line=1.0,
    )


def build_story(doc):
    add_kicker(doc, "01 / Our story")
    add_heading(doc, "Why Delve exists", 1)
    add_para(
        doc,
        "Discovery should not depend on being on social media.",
        size=18,
        bold=True,
        color=INK,
        after=12,
        line=1.05,
    )
    add_para(
        doc,
        (
            "Delve began when founder Tupopila Kadhila wanted to find arts events in "
            "Windhoek, but most of the information was scattered across Instagram. "
            "She did not want to stay on social media simply to know what was happening "
            "in her own city."
        ),
    )
    add_para(
        doc,
        (
            "The same pattern appeared in travel. Many people assume that exploring "
            "Namibia is too expensive, while valuable resident, membership and cardholder "
            "offers - sometimes as much as 50 percent - remain difficult to discover. "
            "Local businesses also promote across disconnected channels and struggle to "
            "turn attention into bookings."
        ),
    )
    add_callout(
        doc,
        "The spark",
        (
            '"I wanted one place where people could find what is happening, see the '
            'deals available to them and explore without depending on social media."'
        ),
    )
    add_callout(
        doc,
        "The problem",
        (
            "Travel, events, offers, transport and local commerce are fragmented across "
            "social media, messaging apps and unrelated booking sites."
        ),
        fill=GOLD,
        accent=GOLD_DARK,
    )
    add_callout(
        doc,
        "Our response",
        (
            "A trusted platform where people discover, compare, book, pay, connect and "
            "review - while businesses manage visibility and demand."
        ),
        fill=MINT,
        accent=MINT_DARK,
    )
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [3.475, 3.475])
    set_table_borders(table, color=LINE, size=5, inside=True)
    for cell, fill, label, text in (
        (
            table.cell(0, 0),
            LILAC,
            "MISSION",
            "Make local experiences, services and savings easier to discover and access.",
        ),
        (
            table.cell(0, 1),
            PALE,
            "VISION",
            "Build a connected travel ecosystem where every city can become a living local network.",
        ),
    ):
        set_cell_shading(cell, fill)
        p = clear_cell(cell)
        set_run(p.add_run(label), size=8, bold=True, color=PURPLE)
        add_para(cell, text, size=10, after=2, line=1.25)


def build_platform(doc):
    add_kicker(doc, "02 / The platform")
    add_heading(doc, "One ecosystem. One account.", 1)
    add_para(
        doc,
        (
            "Delve connects discovery, booking, payment coordination and community "
            "in one mobile-first experience."
        ),
        color=MUTED,
    )
    add_para(
        doc,
        "DISCOVER  >  COMPARE  >  BOOK  >  PAY  >  EXPERIENCE  >  REVIEW",
        size=9,
        bold=True,
        color=PURPLE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=12,
        line=1.0,
    )

    services = [
        (("STAY", "Accommodation"), ("MOVE", "Transport")),
        (("TASTE", "Food and drinks"), ("EXPLORE", "Tours and guides")),
        (("DO", "Activities and events"), ("SHOP", "Local products")),
        (("CONNECT", "Messaging"), ("DELVERS", "Community and stories")),
        (("JOURNEYS", "Routes and memories"), ("TRUST", "Reviews")),
        (("SAVE", "Deals and offers"), ("GROW", "Business profiles")),
    ]
    table = doc.add_table(rows=len(services), cols=2)
    set_table_geometry(table, [3.475, 3.475])
    set_table_borders(table, color=LINE, size=5, inside=True)
    for r_idx, row in enumerate(services):
        for c_idx, (label, detail) in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_shading(cell, PALE if r_idx % 2 == 0 else WHITE)
            p = clear_cell(cell)
            label_run = p.add_run(f"{label}  ")
            set_run(label_run, size=8, bold=True, color=PURPLE)
            detail_run = p.add_run(detail)
            set_run(detail_run, size=9.5, color=INK)

    add_para(doc, "", size=2, after=3, line=1.0)
    audience = doc.add_table(rows=1, cols=2)
    set_table_geometry(audience, [3.475, 3.475])
    set_table_borders(audience, color=LINE, size=5, inside=True)
    traveller = audience.cell(0, 0)
    set_cell_shading(traveller, LILAC)
    clear_cell(traveller)
    add_para(traveller, "FOR TRAVELLERS AND LOCAL EXPLORERS", size=8, bold=True, color=PURPLE, after=5)
    for item in (
        "Find events, stays, transport, food, tours and local products.",
        "Compare options and uncover resident or membership deals.",
        "Message providers, book, review and keep the journey connected.",
    ):
        add_bullet(traveller, item)

    business = audience.cell(0, 1)
    set_cell_shading(business, MINT)
    clear_cell(business)
    add_para(business, "FOR BUSINESSES AND INFORMAL OPERATORS", size=8, bold=True, color=MINT_DARK, after=5)
    for item in (
        "Create a verified profile and publish listings.",
        "Receive inquiries, bookings and orders.",
        "Promote offers, collect reviews and track activity.",
    ):
        add_bullet(business, item)

    add_para(doc, "", size=2, after=3, line=1.0)
    add_callout(
        doc,
        "Fintech role",
        (
            "The current MVP uses simulated payment, payout and dispute workflows. "
            "Live payments are planned through a licensed provider, with provider payouts "
            "released after service completion. Delve will not store raw card details and "
            "will address security, KYC/AML and regulatory requirements before launch."
        ),
        fill=INK,
        accent=WHITE,
    )
    last_table = doc.tables[-1]
    for paragraph in last_table.cell(0, 0).paragraphs[1:]:
        for run in paragraph.runs:
            run.font.color.rgb = rgb(WHITE)


def build_community(doc):
    add_kicker(doc, "03 / Community and memories")
    add_heading(doc, "Everyone is a traveller.", 1)
    add_para(
        doc,
        (
            "Not everyone flies across the world. But everyone goes somewhere, visits "
            "family, explores their city, finds a cafe or discovers something worth sharing."
        ),
        size=13,
        color=PURPLE,
        bold=True,
        after=10,
        line=1.2,
    )
    add_para(
        doc,
        "DELVERS + JOURNEYS TURN EVERYDAY MOVEMENT INTO DISCOVERY.",
        size=8.5,
        bold=True,
        color=INK,
        after=10,
        line=1.0,
    )

    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [3.475, 3.475])
    set_table_borders(table, color=LINE, size=5, inside=True)

    delvers = table.cell(0, 0)
    set_cell_shading(delvers, LILAC)
    clear_cell(delvers)
    add_para(delvers, "DELVERS", size=8, bold=True, color=PURPLE, after=3)
    add_para(delvers, "The living community", size=15, bold=True, after=6, line=1.0)
    add_para(
        delvers,
        (
            "Delvers is where people share real experiences and help one another discover "
            "what is worth seeing, doing and supporting."
        ),
        size=9.5,
        after=6,
        line=1.2,
    )
    for item in (
        "Share photos, videos, stories, travel tips and local discoveries.",
        "Follow creators, destinations, interests and community boards.",
        "Like, save, comment on and discuss useful recommendations.",
        "Move from inspiration to relevant places, offers and services on Delve.",
    ):
        add_bullet(delvers, item)

    journeys = table.cell(0, 1)
    set_cell_shading(journeys, MINT)
    clear_cell(journeys)
    add_para(journeys, "JOURNEYS", size=8, bold=True, color=MINT_DARK, after=3)
    add_para(journeys, "A travel record that helps others", size=15, bold=True, after=6, line=1.0)
    add_para(
        journeys,
        (
            "Journeys lets every traveller document a trip in a useful, structured way - "
            "whether it is across Namibia or across town."
        ),
        size=9.5,
        after=6,
        line=1.2,
    )
    for item in (
        "Record routes, stops, dates, transport modes and total costs.",
        "Add photos, videos, notes, highlights and honest reflections.",
        "Link journey stops to bookable Delve listings and businesses.",
        "Keep a journey public, private or in draft, then share moments to Delvers.",
    ):
        add_bullet(journeys, item)

    add_heading(doc, "How the two work together", 2)
    add_para(
        doc,
        "EXPERIENCE  >  CAPTURE  >  SHARE  >  INSPIRE  >  DISCOVER",
        size=10,
        bold=True,
        color=PURPLE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=4,
        line=1.0,
    )
    add_para(
        doc,
        "Go somewhere  |  Build a Journey  |  Post to Delvers  |  Help another person  |  Start a new journey",
        size=8.5,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=10,
        line=1.0,
    )
    add_callout(
        doc,
        "The Delve difference",
        (
            "Together, Delvers and Journeys make Delve more than a booking platform: "
            "they preserve memories, reveal real costs and turn lived experience into trusted discovery."
        ),
        fill=PALE,
        accent=PURPLE,
    )


def build_market(doc):
    add_kicker(doc, "04 / Market and model")
    add_heading(doc, "Namibia first. Global by design.", 1)
    add_callout(
        doc,
        "Our expansion model",
        (
            "Delve starts by making Namibia's experiences, businesses and deals easier "
            "to discover. The same city-by-city model can then be repeated internationally "
            "while each destination remains locally relevant."
        ),
        fill=PURPLE,
        accent=WHITE,
    )
    for paragraph in doc.tables[-1].cell(0, 0).paragraphs[1:]:
        for run in paragraph.runs:
            run.font.color.rgb = rgb(WHITE)

    add_para(
        doc,
        "INITIAL TARGET USERS",
        size=8,
        bold=True,
        color=INK,
        after=3,
        line=1.0,
    )
    add_para(
        doc,
        "Namibian residents  |  Youth  |  Travellers  |  Informal operators  |  Merchants and MSMEs",
        size=9,
        bold=True,
        color=PURPLE,
        after=10,
        line=1.0,
    )

    columns = doc.add_table(rows=1, cols=3)
    set_table_geometry(columns, [2.317, 2.316, 2.317])
    set_table_borders(columns, color=LINE, size=5, inside=True)
    content = [
        (
            columns.cell(0, 0),
            LILAC,
            PURPLE,
            "WHY NAMIBIA",
            (
                "Events and local offers are fragmented across social channels.",
                "Domestic experiences can feel expensive when discounts are hard to find.",
                "Small businesses need affordable digital visibility and booking tools.",
            ),
        ),
        (
            columns.cell(0, 1),
            GOLD,
            GOLD_DARK,
            "REVENUE MODEL",
            (
                "Booking and transaction commissions.",
                "Business subscriptions and premium tools.",
                "Featured listings and promoted offers.",
                "Advertising, affiliate and travel partnerships.",
            ),
        ),
        (
            columns.cell(0, 2),
            MINT,
            MINT_DARK,
            "CURRENT PROGRESS",
            (
                "Functional mobile-first MVP.",
                "Traveller, provider and admin interfaces.",
                "Booking, verification and moderation workflows.",
                "Simulated payment, payout and dispute systems.",
            ),
        ),
    ]
    for cell, fill, accent, title, items in content:
        set_cell_shading(cell, fill)
        clear_cell(cell)
        add_para(cell, title, size=8, bold=True, color=accent, after=6)
        for item in items:
            add_bullet(cell, item)

    add_heading(doc, "Next phase", 2)
    for item in (
        "Complete security and performance testing.",
        "Integrate a licensed payment provider.",
        "Run a controlled pilot with Namibian events, tourism and MSME partners.",
        "Use pilot data to validate pricing, demand and the addressable market.",
    ):
        add_numbered(doc, item)
    add_para(
        doc,
        (
            "Market size is still being formally assessed; the pilot will provide "
            "evidence for a reliable estimate."
        ),
        size=8.5,
        italic=True,
        color=MUTED,
        after=0,
        line=1.0,
    )


def build_team(doc):
    add_kicker(doc, "05 / Leadership")
    add_heading(doc, "The team building Delve", 1)
    add_para(
        doc,
        (
            "A founder-led team combining product development, operations, financial "
            "management and strategic guidance."
        ),
        color=MUTED,
        after=10,
    )

    team = doc.add_table(rows=1, cols=3)
    set_table_geometry(team, [2.317, 2.316, 2.317])
    set_table_borders(team, color=LINE, size=5, inside=True)
    members = [
        (
            "tupopila.jpg",
            "Tupopila Kadhila",
            "FOUNDER, CEO AND PRODUCT LEAD",
            "Software developer building Delve from the ground up and leading the product vision.",
        ),
        (
            "michael.jpg",
            "Michael Kadhila",
            "CO-FOUNDER AND HEAD OF OPERATIONS",
            "Leads operations, partnerships, financial management and the foundation needed to scale.",
        ),
        (
            "slysken.jpg",
            "Slysken Kakuva",
            "CO-FOUNDER AND BOARD OBSERVER",
            "Provides strategic guidance, industry insight and independent feedback for long-term growth.",
        ),
    ]
    for idx, (image_name, name, role, description) in enumerate(members):
        cell = team.cell(0, idx)
        set_cell_shading(cell, WHITE)
        p = clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(ASSET_DIR / image_name), width=Inches(1.95))
        add_para(cell, name, size=11.5, bold=True, after=3, line=1.0)
        add_para(cell, role, size=7.2, bold=True, color=PURPLE, after=6, line=1.0)
        add_para(cell, description, size=9, after=2, line=1.2)

    add_para(doc, "", size=2, after=3, line=1.0)
    governance = doc.add_table(rows=1, cols=2)
    set_table_geometry(governance, [3.475, 3.475])
    set_table_borders(governance, color=LINE, size=5, inside=True)

    gov = governance.cell(0, 0)
    set_cell_shading(gov, LILAC)
    clear_cell(gov)
    add_para(gov, "GOVERNANCE", size=8, bold=True, color=PURPLE, after=6)
    for item in (
        "Founder-led governance with strategic observer support.",
        "The CEO and Head of Operations review material expenditure.",
        "Formal policies, reporting controls and an advisory structure will be strengthened as Delve grows.",
    ):
        add_bullet(gov, item)

    priorities = governance.cell(0, 1)
    set_cell_shading(priorities, MINT)
    clear_cell(priorities)
    add_para(priorities, "12-MONTH PRIORITIES", size=8, bold=True, color=MINT_DARK, after=6)
    for item in (
        "Bring additional development capacity into the team.",
        "Complete payment, security and compliance readiness.",
        "Onboard pilot businesses and event organisers.",
        "Measure user demand, bookings and partner outcomes.",
    ):
        add_numbered(priorities, item)

    add_para(doc, "", size=2, after=3, line=1.0)
    contact = doc.add_table(rows=1, cols=1)
    set_table_geometry(contact, [6.95])
    remove_table_borders(contact)
    cell = contact.cell(0, 0)
    set_cell_shading(cell, PURPLE)
    clear_cell(cell)
    add_para(
        cell,
        "Let us build a more connected way to explore.",
        size=16,
        bold=True,
        color=WHITE,
        after=8,
        line=1.05,
    )
    add_para(
        cell,
        "DELVE  |  Windhoek, Namibia  |  delveworldwide@gmail.com  |  +264 81 764 9719",
        size=8.5,
        bold=True,
        color=WHITE,
        after=0,
        line=1.0,
    )


def build_docx():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    if not LOGO_PATH.exists():
        raise FileNotFoundError(LOGO_PATH)
    for image_name in ("tupopila.jpg", "michael.jpg", "slysken.jpg"):
        if not (ASSET_DIR / image_name).exists():
            raise FileNotFoundError(ASSET_DIR / image_name)

    doc = Document()
    set_styles(doc)
    configure_page(doc)
    set_page_furniture(doc)
    props = doc.core_properties
    props.title = "Delve Company Profile 2026"
    props.subject = "Company profile for Delve"
    props.author = "Delve"
    props.keywords = "Delve, Namibia, travel, events, marketplace, fintech"

    build_cover(doc)
    add_page_break(doc)
    build_story(doc)
    add_page_break(doc)
    build_platform(doc)
    add_page_break(doc)
    build_community(doc)
    add_page_break(doc)
    build_market(doc)
    add_page_break(doc)
    build_team(doc)

    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build_docx()
