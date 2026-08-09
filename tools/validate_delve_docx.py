from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree


path = Path("output/docx/Delve_Company_Profile_2026.docx")
archive = ZipFile(path)
assert archive.testzip() is None

document = Document(path)
all_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
all_text += "\n" + "\n".join(
    cell.text
    for table in document.tables
    for row in table.rows
    for cell in row.cells
)

root = etree.fromstring(archive.read("word/document.xml"))
namespace = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
}
page_breaks = root.xpath('.//w:br[@w:type="page"]', namespaces=namespace)

print("VALID_ZIP", True)
print("PAGE_BREAKS", len(page_breaks))
print("TABLES", len(document.tables))
print("IMAGES", len(document.inline_shapes))
print(
    "MOTTO_OK",
    "Everyone is a traveller" in all_text
    and "Everybody is a traveller" not in all_text,
)
print("DELVERS_JOURNEYS_OK", "DELVERS" in all_text and "JOURNEYS" in all_text)
print(
    "PAGE_CM",
    round(document.sections[0].page_width.cm, 2),
    round(document.sections[0].page_height.cm, 2),
)

assert len(page_breaks) == 5
assert len(document.inline_shapes) == 4
assert "Everyone is a traveller" in all_text
assert "Everybody is a traveller" not in all_text
assert "DELVERS" in all_text
assert "JOURNEYS" in all_text
